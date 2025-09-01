"""
src/ingest.py
Extract text from uploaded files for the simplified Idea2Paper pipeline.

Supported:
- PDF: prefers PyMuPDF (pymupdf) for cleaner text; falls back to pypdf.
- DOCX: prefers python-docx (handles paragraphs + tables); falls back to docx2txt.
- TXT / MD: UTF-8 with latin-1 fallback.
- CSV: pandas -> compact CSV string (header + up to N rows).

Public API:
- extract_text(file_bytes: bytes, filename: str) -> str
- extract_many(files: List[UploadedFile], min_chars: int = 200) -> List[Tuple[str, str]]

Notes:
- Image-only/encrypted PDFs will return empty text (no OCR).
- Very large inputs are the caller's responsibility to trim/limit downstream.
"""

from __future__ import annotations

import io
import re
from typing import List, Tuple

# Optional deps (graceful fallbacks)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None  # type: ignore

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None  # type: ignore

try:
    from docx import Document  # python-docx
except Exception:
    Document = None  # type: ignore

try:
    import docx2txt  # fallback
except Exception:
    docx2txt = None  # type: ignore

try:
    import pandas as pd
except Exception:
    pd = None  # type: ignore


SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".csv"}


# ------------------------
# Utilities
# ------------------------

def _normalize_ws(text: str) -> str:
    if not text:
        return ""
    # Remove non-printables, collapse whitespace/newlines
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _guess_ext(name: str) -> str:
    if not name:
        return ""
    s = name.lower().strip()
    for ext in SUPPORTED_EXTS:
        if s.endswith(ext):
            return ext
    m = re.search(r"(\.[a-z0-9]{1,6})$", s)
    return m.group(1) if m else ""


# ------------------------
# PDF extraction
# ------------------------

def _extract_text_from_pdf_pymupdf(file_bytes: bytes) -> str:
    if not fitz:
        return ""
    try:
        text_parts: List[str] = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text_parts.append(page.get_text("text") or "")
        return _normalize_ws("\n".join(text_parts))
    except Exception:
        return ""


def _extract_text_from_pdf_pypdf(file_bytes: bytes) -> str:
    if not PdfReader:
        return ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts: List[str] = []
        for page in getattr(reader, "pages", []):
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        return _normalize_ws("\n".join(parts))
    except Exception:
        return ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Prefer PyMuPDF, fallback to PyPDF.
    """
    txt = _extract_text_from_pdf_pymupdf(file_bytes)
    if txt:
        return txt
    return _extract_text_from_pdf_pypdf(file_bytes)


# ------------------------
# DOCX extraction
# ------------------------

def _extract_text_from_docx_python_docx(file_bytes: bytes) -> str:
    if not Document:
        return ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: List[str] = []
        # Paragraphs
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        # Tables
        for tbl in doc.tables:
            for row in tbl.rows:
                row_text = "\t".join((cell.text or "").strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return _normalize_ws("\n".join(parts))
    except Exception:
        return ""


def _extract_text_from_docx_docx2txt(file_bytes: bytes) -> str:
    if not docx2txt:
        return ""
    try:
        # docx2txt needs a file path; use a temp buffer on disk only if needed
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
            tmp.write(file_bytes)
            tmp.flush()
            txt = docx2txt.process(tmp.name) or ""
            return _normalize_ws(txt)
    except Exception:
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Prefer python-docx (captures tables), fallback to docx2txt.
    """
    txt = _extract_text_from_docx_python_docx(file_bytes)
    if txt:
        return txt
    return _extract_text_from_docx_docx2txt(file_bytes)


# ------------------------
# Text / Markdown / CSV
# ------------------------

def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        txt = file_bytes.decode("utf-8", errors="ignore")
        if not txt.strip():
            txt = file_bytes.decode("latin-1", errors="ignore")
        return _normalize_ws(txt)
    except Exception:
        return ""


def extract_text_from_md(file_bytes: bytes) -> str:
    # Treat Markdown as plain text here (downstream summarizer doesn't need MD semantics)
    return extract_text_from_txt(file_bytes)


def extract_text_from_csv(file_bytes: bytes, max_rows: int = 200) -> str:
    if pd is None:
        return ""
    try:
        df = pd.read_csv(io.BytesIO(file_bytes))
    except Exception:
        try:
            df = pd.read_csv(io.BytesIO(file_bytes), encoding="latin-1")
        except Exception:
            return ""
    if len(df) > max_rows:
        df = df.head(max_rows)
    try:
        return df.to_csv(index=False)
    except Exception:
        return df.astype(str).to_csv(index=False)


# ------------------------
# Public API
# ------------------------

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch by extension and return best-effort text content.
    """
    ext = _guess_ext(filename)
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if ext == ".docx":
        return extract_text_from_docx(file_bytes)
    if ext == ".txt":
        return extract_text_from_txt(file_bytes)
    if ext == ".md":
        return extract_text_from_md(file_bytes)
    if ext == ".csv":
        return extract_text_from_csv(file_bytes)
    # Unknown: try utf-8 as a last resort
    try:
        return _normalize_ws(file_bytes.decode("utf-8", errors="ignore"))
    except Exception:
        return ""


def extract_many(files: List, min_chars: int = 200) -> List[Tuple[str, str]]:
    """
    Streamlit `st.file_uploader` → list of (filename, extracted_text).
    Skips entries whose extracted text is shorter than `min_chars`.
    """
    out: List[Tuple[str, str]] = []
    for f in files or []:
        try:
            name = getattr(f, "name", "upload")
            # Streamlit UploadedFile supports .read() or .getvalue()
            data = f.read() if hasattr(f, "read") else getattr(f, "getvalue", lambda: b"")()
            if not data:
                continue
            txt = extract_text(data, name)
            if txt and len(txt) >= int(min_chars):
                out.append((name, txt))
        except Exception:
            # Skip this file on any error
            continue
    return out


__all__ = [
    "extract_text",
    "extract_many",
    "extract_text_from_pdf",
    "extract_text_from_docx",
    "extract_text_from_txt",
    "extract_text_from_md",
    "extract_text_from_csv",
]
